import os
import subprocess
from pathlib import Path
from typing import Optional, List
from docx import Document as DocxDocument
import pdfplumber

class Document:
    def __init__(self, docx_path: Path):
        if not docx_path.exists() or docx_path.suffix.lower() != ".docx":
            raise ValueError("Document must be a valid .docx file")
        self.folder = docx_path.parent
        self.base_name = docx_path.stem
        self.original_docx_path = docx_path
        self.original_pdf_path = self.folder / f"{self.base_name}.pdf"
        self.translated_docx_path = self.folder / f"{self.base_name}_translated.docx"
        self.translated_pdf_path = self.folder / f"{self.base_name}_translated.pdf"
        self.docx_path_to_use: Optional[Path] = None
        self.pdf_path_to_use: Optional[Path] = None

    @classmethod
    def from_file(cls, file_path: Path) -> "Document":
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        if file_path.suffix.lower() == ".docx":
            return cls(file_path)
        elif file_path.suffix.lower() == ".pdf":
            docx_path = cls.convert_pdf_to_docx(file_path)
            return cls(docx_path)
        else:
            raise ValueError("Unsupported file type. Only .docx or .pdf are allowed.")

    @staticmethod
    def convert_pdf_to_docx(pdf_path: Path) -> Path:
        if not pdf_path.exists():
            raise FileNotFoundError("PDF file does not exist.")
        docx_path = pdf_path.with_suffix(".docx")
        doc = DocxDocument()
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
        doc.save(str(docx_path))
        return docx_path

    @staticmethod
    def convert_docx_to_pdf(docx_path: Path) -> Path:
        if not docx_path.exists():
            raise FileNotFoundError("DOCX file does not exist.")
        pdf_path = docx_path.with_suffix(".pdf")
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", str(docx_path),
            "--outdir", str(docx_path.parent)
        ], check=True)
        return pdf_path

    def ensure_pdf_exists(self):
        if not self.original_pdf_path.exists():
            self.original_pdf_path = self.convert_docx_to_pdf(self.original_docx_path)

    def extract_text(self) -> str:
        doc = DocxDocument(str(self.original_docx_path))
        return "".join([para.text for para in doc.paragraphs]).strip()

    def get_paragraphs(self) -> List[str]:
        doc = DocxDocument(str(self.original_docx_path))
        return [para.text for para in doc.paragraphs if para.text.strip()]

    def save_translated(self, translated_texts: List[str]):
        doc = DocxDocument()
        for text in translated_texts:
            doc.add_paragraph(text)
        doc.save(str(self.translated_docx_path))
        self.convert_docx_to_pdf(self.translated_docx_path)

    def set_paths_to_use(self, translated: bool):
        if translated:
            self.docx_path_to_use = self.translated_docx_path
            self.pdf_path_to_use = self.translated_pdf_path
        else:
            self.docx_path_to_use = self.original_docx_path
            self.pdf_path_to_use = self.original_pdf_path
